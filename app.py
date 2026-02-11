from flask import Flask, render_template, request, send_file, abort
from flask import redirect
import pandas as pd
import os
import uuid

from docx import Document
from reportlab.lib.pagesizes import A4, landscape
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet
from openpyxl.utils import get_column_letter

app = Flask(__name__)

# ===============================
# PATH
# ===============================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads", "excel")
OUTPUT_FOLDER = os.path.join(BASE_DIR, "uploads", "hasil_excel")

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

# ===============================
# SESSION CACHE (AMAN PUBLIK)
# ===============================
DATA_CACHE = {}

# ===============================
# FORMAT
# ===============================
def format_nominal(val):
    if pd.isna(val):
        return ""
    try:
        if isinstance(val, (int, float)):
            return f"{int(val):,}".replace(",", ".")
    except:
        pass
    return str(val)

def format_datetime(val):
    if isinstance(val, pd.Timestamp):
        return val.strftime("%d-%m-%Y")
    return val

# ===============================
# INDEX
# ===============================
@app.route("/")
def index():
    return render_template(
        "index.html",
        projects=[],
        message="Silakan upload file Excel"
    )

# ===============================
# UPLOAD
# ===============================
@app.route("/upload", methods=["POST"])
def upload():
    file = request.files.get("file")
    if not file:
        return "File tidak ditemukan", 400

    session_id = str(uuid.uuid4())
    save_path = os.path.join(UPLOAD_FOLDER, f"{session_id}_{file.filename}")
    file.save(save_path)

    # Cari header otomatis
    df_raw = pd.read_excel(save_path, header=None)
    header_row = next(
        (i for i, r in df_raw.iterrows() if r.notna().sum() >= 2),
        None
    )
    if header_row is None:
        return "Header tidak ditemukan", 400

    df = pd.read_excel(save_path, header=header_row)

    # Bersihkan data
    df = df.replace(r'^\s*$', pd.NA, regex=True)
    df = df.dropna(axis=1, how="all")
    df = df.loc[:, ~df.columns.str.contains("^Unnamed", na=False)]
    df = df.reset_index(drop=True)

    # Format tanggal
    for col in df.columns:
        if pd.api.types.is_datetime64_any_dtype(df[col]):
            df[col] = df[col].apply(format_datetime)

    # Cari kolom kode kegiatan
    kolom_kode = next(
        (c for c in df.columns
         if "kode kegiatan" in str(c).lower()
         or "kode keg" in str(c).lower()),
        None
    )
    if not kolom_kode:
        return "Kolom Kode Kegiatan tidak ditemukan", 400

    DATA_CACHE[session_id] = {
        "df": df,
        "kode_col": kolom_kode
    }

    kode_list = (
        df[kolom_kode]
        .dropna()
        .astype(str)
        .unique()
        .tolist()
    )

    return render_template(
        "index.html",
        projects=kode_list,
        session_id=session_id,
        message=None
    )

# ===============================
# WORD
# ===============================
def buat_word(data, filename):
    path = os.path.join(OUTPUT_FOLDER, filename)
    doc = Document()
    doc.add_heading("Data Kode Kegiatan", level=1)

    table = doc.add_table(rows=1, cols=len(data.columns))
    table.style = "Table Grid"

    for i, col in enumerate(data.columns):
        table.rows[0].cells[i].text = str(col)

    for _, row in data.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = format_nominal(val)

    doc.save(path)
    return path

# ===============================
# PDF
# ===============================
def buat_pdf(data, filename):
    path = os.path.join(OUTPUT_FOLDER, filename)

    styles = getSampleStyleSheet()
    style = styles["Normal"]
    style.fontSize = 7
    style.leading = 9

    table_data = [[Paragraph(str(c), style) for c in data.columns]]
    for _, row in data.iterrows():
        table_data.append(
            [Paragraph(format_nominal(v), style) for v in row]
        )

    table = Table(table_data, repeatRows=1)
    table.setStyle(TableStyle([
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
    ]))

    doc = SimpleDocTemplate(
        path,
        pagesize=landscape(A4),
        rightMargin=15,
        leftMargin=15,
        topMargin=15,
        bottomMargin=15
    )

    doc.build([table])
    return path

# ===============================
# DETAIL
# ===============================
@app.route("/detail/<session_id>/<kode>")
def detail(session_id, kode):
    data_pack = DATA_CACHE.get(session_id)
    if not data_pack:
        return "Session tidak valid", 400

    df = data_pack["df"]
    kolom_kode = data_pack["kode_col"]

    data = df[df[kolom_kode].astype(str) == kode]
    if data.empty:
        return "Data tidak ditemukan", 404

    excel_name = f"{session_id}_{kode}.xlsx"
    word_name = f"{session_id}_{kode}.docx"
    pdf_name = f"{session_id}_{kode}.pdf"

    # Excel
    excel_path = os.path.join(OUTPUT_FOLDER, excel_name)
    with pd.ExcelWriter(excel_path, engine="openpyxl") as writer:
        data.to_excel(writer, index=False)
        ws = writer.sheets["Sheet1"]
        for i, col in enumerate(data.columns, 1):
            max_len = max(len(str(col)), 12)
            ws.column_dimensions[get_column_letter(i)].width = max_len + 2

    # Word & PDF
    buat_word(data, word_name)
    buat_pdf(data, pdf_name)

    return render_template(
        "detail.html",
        kode=kode,
        session_id=session_id,
        table=data.apply(lambda c: c.map(format_nominal)).to_html(index=False),
        excel=excel_name,
        word=word_name,
        pdf=pdf_name
    )

# ===============================
# OPEN EXCEL (TANPA DOWNLOAD)
# ===============================
@app.route("/open-excel/<filename>")
def open_excel(filename):
    path = os.path.join(OUTPUT_FOLDER, filename)
    if not os.path.exists(path):
        abort(404)

    return send_file(
        path,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        as_attachment=False
    )

# ===============================
# EXCEL ONLINE (AUTO BUKA MICROSOFT)
# ===============================
@app.route("/excel-online/<filename>")
def excel_online(filename):
    path = os.path.join(OUTPUT_FOLDER, filename)
    if not os.path.exists(path):
        abort(404)

    # URL publik file di Railway
    file_url = request.url_root.rstrip("/") + "/open-excel/" + filename

    # Redirect ke Microsoft Office Web Viewer
    office_url = (
        "https://view.officeapps.live.com/op/view.aspx?src="
        + file_url
    )

    return redirect(office_url)

@app.route("/preview-excel/<filename>")
def preview_excel(filename):
    path = os.path.join(OUTPUT_FOLDER, filename)
    if not os.path.exists(path):
        abort(404)

    df = pd.read_excel(path)

    table = df.apply(lambda c: c.map(format_nominal)).to_html(
        index=False,
        classes="excel-table"
    )

    return render_template(
        "preview.html",
        table=table,
        filename=filename
    )
    
# ===============================
# DOWNLOAD
# ===============================
@app.route("/download/<filename>")
def download(filename):
    path = os.path.join(OUTPUT_FOLDER, filename)
    if not os.path.exists(path):
        abort(404)
    return send_file(path, as_attachment=True)

# ===============================
# RUN
# ===============================
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    app.run(host="0.0.0.0", port=port)

